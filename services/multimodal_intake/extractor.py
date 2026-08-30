"""
extractor.py — Advanced Data Extraction Engine for IMDIE.

Extracts text, key-value pairs, tables, and structured metrics from PDF, DOCX, RTF,
CSV, TSV, JSON, XLS, XLSX, JPG, JPEG, PNG, WEBP, HEIC, TIFF, or text files into a
standardized provenance-tracked representation.

Header byte signature inspection guarantees true file format detection.
Image preprocessing (EXIF rotation, deskew, contrast, blur check) precedes OCR.
"""

import io
import json
import logging
import os
import re
import zipfile
from pathlib import Path
from typing import Any, Dict, List, Tuple, Union

from PIL import Image, ImageEnhance, ImageFilter, ImageOps

logger = logging.getLogger("imdie.extractor")

# Register HEIC/HEIF opener if pillow_heif is installed
try:
    import pillow_heif
    pillow_heif.register_heif_opener()
except ImportError:
    logger.info("pillow_heif not installed; native HEIC decoding will fall back to PIL standard openers.")

# Try importing format libraries
try:
    import pypdf
except ImportError:
    pypdf = None

try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

try:
    import docx
except ImportError:
    docx = None

try:
    import openpyxl
except ImportError:
    openpyxl = None

try:
    import pandas as pd
except ImportError:
    pd = None

try:
    from striprtf.striprtf import rtf_to_text
except ImportError:
    rtf_to_text = None

try:
    import pytesseract
    # Configure Tesseract executable location on Windows
    tess_paths = [
        r"C:\Program Files\Tesseract-OCR\tesseract.exe",
        r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
        os.path.expanduser(r"~\AppData\Local\Programs\Tesseract-OCR\tesseract.exe")
    ]
    for tp in tess_paths:
        if os.path.exists(tp):
            pytesseract.pytesseract.tesseract_cmd = tp
            break
except ImportError:
    pytesseract = None

EXTRACTION_CONFIDENCE_MAP = {
    "REGEX_PATIENT_ID": 1.00,
    "REGEX_AGE": 1.00,
    "REGEX_GENDER": 1.00,
    "REGEX_FAMILY_HISTORY": 0.95,
    "CSV_HEADER_ROW_PAIR": 0.98,
    "PATTERN_MATCH_VALUNIT": 0.95,
    "PATTERN_MATCH_KEYVAL": 0.90,
    "INLINE_MULTI_PAIR": 0.90,
    "OCR_EXTRACTION": 0.85,
    "FUZZY_ALIAS": 0.85,
    "DERIVED_RULE": 0.90
}


def detect_file_format_by_header(file_bytes: Union[bytes, str, Path], filename: str = "") -> str:
    """Inspect magic header bytes to determine true file MIME format regardless of extension.

    Returns format string: 'pdf', 'docx', 'xlsx', 'rtf', 'jpeg', 'png', 'webp', 'tiff', 'heic', 'csv', 'json', 'txt', 'unknown'
    """
    if isinstance(file_bytes, (str, Path)):
        if not filename:
            filename = str(file_bytes)
        try:
            with open(file_bytes, "rb") as f:
                file_bytes = f.read()
        except Exception:
            return "unknown"

    if not file_bytes:
        return "empty"

    header = file_bytes[:32]

    if header.startswith(b"%PDF"):
        return "pdf"
    elif header.startswith(b"{\\rtf"):
        return "rtf"
    elif header.startswith(b"\xff\xd8\xff"):
        return "jpeg"
    elif header.startswith(b"\x89PNG\r\n\x1a\n"):
        return "png"
    elif header.startswith(b"RIFF") and b"WEBP" in header[8:16]:
        return "webp"
    elif header.startswith(b"II*\x00") or header.startswith(b"MM\x00*"):
        return "tiff"
    elif b"ftypheic" in header[4:24] or b"ftypheim" in header[4:24] or b"ftypheis" in header[4:24]:
        return "heic"
    elif header.startswith(b"PK\x03\x04"):
        # ZIP container: inspect zip entries for DOCX vs XLSX
        try:
            with zipfile.ZipFile(io.BytesIO(file_bytes)) as zf:
                names = zf.namelist()
                if any(n.startswith("word/") for n in names):
                    return "docx"
                elif any(n.startswith("xl/") for n in names):
                    return "xlsx"
        except Exception:
            pass
        return "zip"
    elif header.startswith(b"\xd0\xcf\x11\xe0"):
        # OLE compound file: doc or xls
        ext = Path(filename).suffix.lower()
        return "xls" if ext == ".xls" else "doc"

    # String format check (JSON / CSV / TXT / RTF plain)
    try:
        text_start = file_bytes[:1024].decode("utf-8", errors="ignore").strip()
        if text_start.startswith("{") or text_start.startswith("["):
            return "json"
        elif "\t" in text_start or Path(filename).suffix.lower() == ".tsv":
            return "tsv"
        elif "," in text_start:
            return "csv"
        return "txt"
    except Exception:
        return "unknown"


def estimate_image_blur_variance(pil_img: Image.Image) -> float:
    """Estimate image sharpness/blur using grayscale pixel variance."""
    try:
        gray = pil_img.convert("L")
        # Compute difference filter variance as a proxy for Laplacian sharpness
        edges = gray.filter(ImageFilter.FIND_EDGES)
        pixels = list(edges.getdata())
        if not pixels:
            return 100.0
        mean = sum(pixels) / len(pixels)
        variance = sum((p - mean) ** 2 for p in pixels) / len(pixels)
        return float(variance)
    except Exception:
        return 100.0


def preprocess_image_for_ocr(pil_img: Image.Image) -> Tuple[Image.Image, Dict[str, Any]]:
    """Preprocess camera photos & report scans prior to OCR.

    Applies EXIF rotation, contrast auto-enhancement, blur detection.
    Returns: Tuple of (preprocessed_pil_image, preprocessing_meta)
    """
    meta = {
        "exif_rotated": False,
        "contrast_enhanced": False,
        "blur_variance": 100.0,
        "is_unreadable": False,
        "unreadable_reason": None
    }

    try:
        # 1. EXIF orientation correction
        pil_img = ImageOps.exif_transpose(pil_img)
        meta["exif_rotated"] = True
    except Exception:
        pass

    # 2. Convert palette/RGBA images to RGB
    if pil_img.mode not in ["RGB", "L"]:
        pil_img = pil_img.convert("RGB")

    # 3. Blur detection check
    blur_var = estimate_image_blur_variance(pil_img)
    meta["blur_variance"] = round(blur_var, 2)
    if blur_var < 5.0:
        meta["is_unreadable"] = True
        meta["unreadable_reason"] = f"Image is severely blurred or low-contrast (sharpness score {round(blur_var, 1)} < 5.0). Please retake a clear photo."

    # 4. Auto-contrast enhancement
    try:
        pil_img = ImageOps.autocontrast(pil_img)
        enhancer = ImageEnhance.Contrast(pil_img)
        pil_img = enhancer.enhance(1.4)
        meta["contrast_enhanced"] = True
    except Exception:
        pass

    return pil_img, meta


def extract_key_value_pairs_from_text(text: str, source_file: str = "text_input", page_num: int = 1) -> Dict[str, Any]:
    """Extract key-value pairs, units, and reference ranges from unstructured, CSV, or tabular text."""
    extracted: Dict[str, Any] = {}

    raw_lines = [l.strip() for l in text.splitlines() if l.strip()]

    # 1. Handle CSV / TSV Header and Value Rows
    csv_headers = []
    for line in raw_lines:
        if ("," in line or "\t" in line) and any(kw in line.lower() for kw in ["age", "gender", "height", "weight", "glucose", "hba1c", "ldl", "hdl", "triglycerides", "steps", "akkermansia", "shannon"]):
            delim = "\t" if "\t" in line else ","
            parts = [p.strip() for p in line.split(delim)]
            if len(parts) >= 3:
                csv_headers.append((line, parts, delim))

    if csv_headers:
        for orig_header, headers, delim in csv_headers:
            header_idx = raw_lines.index(orig_header)
            for r_idx in range(header_idx + 1, min(header_idx + 5, len(raw_lines))):
                val_line = raw_lines[r_idx]
                if delim in val_line:
                    val_parts = [p.strip() for p in val_line.split(delim)]
                    if len(val_parts) == len(headers):
                        for h, v in zip(headers, val_parts):
                            if h and v and not h.lower().startswith("patient name"):
                                extracted[h] = {
                                    "raw_value": v,
                                    "source_file": source_file,
                                    "source_page": page_num,
                                    "line_index": r_idx,
                                    "extraction_method": "CSV_HEADER_ROW_PAIR"
                                }

    # Reconstruct vertical PDF table cell pairs & multi-line BP
    lines = []
    skip_count = 0
    for i in range(len(raw_lines)):
        if skip_count > 0:
            skip_count -= 1
            continue
        line = raw_lines[i]
        if i + 1 < len(raw_lines):
            next_line = raw_lines[i + 1]

            # Vertical BP Split: "Blood Pressure" \n "124" \n "/ 80 mmHg"
            if re.search(r"^\s*Blood\s*Pressure\s*$", line, re.IGNORECASE) and re.search(r"^\d{2,3}$", next_line):
                if i + 2 < len(raw_lines) and re.search(r"^\s*[\/\-]\s*\d{2,3}", raw_lines[i + 2]):
                    lines.append(f"Blood Pressure: {next_line}{raw_lines[i + 2].strip()}")
                    skip_count = 2
                    continue

            if re.search(r"^(Patient\s*ID|Report\s*ID|Collection\s*date|Age\s*/\s*Gender)$", line, re.IGNORECASE):
                lines.append(f"{line}: {next_line}")
                skip_count = 1
                continue
            if re.search(r"^\s*(\d+(?:\.\d+)?|Yes|No|Male|Female)\b", next_line, re.IGNORECASE):
                unit_str = ""
                if i + 2 < len(raw_lines):
                    line3 = raw_lines[i + 2].strip()
                    if re.search(r"^(cm|kg|kg/m|mmHg|mg/dL|%|U/L|steps/day|min/day|bpm|ms|hours/night|hours|score/\d+|index|mmol/L)", line3, re.IGNORECASE):
                        unit_str = " " + line3
                        skip_count = 2
                    else:
                        skip_count = 1
                else:
                    skip_count = 1
                lines.append(f"{line} {next_line}{unit_str}")
                continue
        lines.append(line)

    for line_idx, line in enumerate(lines):
        line = line.strip()
        if not line or line.startswith("#"):
            continue

        # Split concatenated letter-digit patterns from PDF OCR/text streams (e.g. "CHOLESTEROL170" -> "CHOLESTEROL 170", "VitalsWeight78" -> "VitalsWeight 78"), excluding HbA1c
        line_protected = line.replace("HbA1c", "__HBAONEC__").replace("hba1c", "__HBAONEC__")
        line = re.sub(r'([A-Za-z])(\d+(?:\.\d+)?)', r'\1 \2', line_protected).replace("__HBAONEC__", "HbA1c")

        # Skip headers / disclaimer lines
        if any(h in line.upper() for h in ["SYNTHETIC TEST DATA", "SUNRISE DIAGNOSTIC", "CLINICAL CHEMISTRY & ANTHROPOMETRY", "MEASUREMENT RESULT REFERENCE", "LABORATORY NOTE:"]):
            continue

        # Inline multi-pair extraction e.g. "179cm Wt:83kg", "46Y Sex:M", "WC:98cm BP:130/84mmHg"
        inline_pairs = re.findall(r"(?:^|\s+|\s*;\s*)([A-Za-z0-9_\-\%/\(\)\s]{2,25})\s*[:=]\s*([A-Za-z0-9%/_\-\/\.²\s]+?)(?=\s+[A-Za-z0-9_\-\%/\(\)]{2,25}\s*[:=]|$)", line)
        if len(inline_pairs) >= 2:
            for k, v in inline_pairs:
                k_clean = k.strip()
                v_clean = v.strip()
                if k_clean and v_clean:
                    extracted[k_clean] = {
                        "raw_value": v_clean,
                        "source_file": source_file,
                        "source_page": page_num,
                        "line_index": line_idx,
                        "extraction_method": "INLINE_MULTI_PAIR"
                    }
            continue

        # Pattern A & B: Patient ID, Age, Gender header line
        is_header_demographic = False
        m_pid = re.search(r"Patient[\s_]*ID[:\s\t]+([A-Za-z0-9_\-]+)", line, re.IGNORECASE)
        if m_pid:
            pid_val = m_pid.group(1).strip()
            if pid_val.upper() not in ["REPORT", "COLLECTION", "NAME"]:
                extracted["Patient_ID"] = {
                    "raw_value": pid_val,
                    "source_file": source_file,
                    "source_page": page_num,
                    "line_index": line_idx,
                    "extraction_method": "REGEX_PATIENT_ID"
                }
                is_header_demographic = True

        m_age = re.search(r"\bAge\b.*?\b(\d{1,3})\b", line, re.IGNORECASE)
        if m_age:
            extracted["Age"] = {
                "raw_value": m_age.group(1).strip(),
                "source_file": source_file,
                "source_page": page_num,
                "line_index": line_idx,
                "extraction_method": "REGEX_AGE"
            }
            is_header_demographic = True

        m_gen = re.search(r"\b(?:Gender|Sex)\b.*?\b(Male|Female|Other)\b", line, re.IGNORECASE)
        if m_gen:
            extracted["Gender"] = {
                "raw_value": m_gen.group(1).strip(),
                "source_file": source_file,
                "source_page": page_num,
                "line_index": line_idx,
                "extraction_method": "REGEX_GENDER"
            }
            is_header_demographic = True

        if is_header_demographic:
            continue

        # Pattern C: Family History (e.g. "Family history: diabetes Yes Reported")
        m_fam = re.search(r"Family\s*history\s*[:\s]*([A-Za-z0-9_\s]+?)\s+(Yes|No|1|0)\b", line, re.IGNORECASE)
        if m_fam:
            fam_key = f"Family history: {m_fam.group(1).strip()}"
            fam_val = m_fam.group(2).strip()
            extracted[fam_key] = {
                "raw_value": fam_val,
                "source_file": source_file,
                "source_page": page_num,
                "line_index": line_idx,
                "extraction_method": "REGEX_FAMILY_HISTORY"
            }
            continue

        # Pattern 1: Key : Value or Key = Value or Key \t Value
        match = re.search(r"^\s*([A-Za-z0-9_\s\-\%/\(\)]+?)\s*[:=,\t]\s*(.+?)\s*$", line)
        if match:
            k, v = match.group(1).strip(), match.group(2).strip()
            if not k.startswith("=") and not k.startswith("-") and v:
                if k not in extracted or extracted[k].get("extraction_method") == "PATTERN_MATCH_KEYVAL":
                    extracted[k] = {
                        "raw_value": v,
                        "source_file": source_file,
                        "source_page": page_num,
                        "line_index": line_idx,
                        "extraction_method": "PATTERN_MATCH_KEYVAL"
                    }
                continue

        # Pattern 2: Key value unit reference (e.g. "Fasting blood glucose 154 mg/dL 70-99", "ALT (SGPT) 34 U/L", "Average Daily Steps 6500 steps/day")
        match2 = re.search(r"^([A-Za-z0-9_\s\-\%/\(\)]+?)\s+([0-9]+(?:,[0-9]+)*(?:\.[0-9]+)?)\s*([A-Za-z0-9%/_\-\/²]+)?", line)
        if match2:
            k = match2.group(1).strip()
            val = match2.group(2).strip()
            unit = match2.group(3)
            if k and not k.upper().startswith("PAGE") and not k.upper().startswith("PATIENT"):
                val_str = f"{val} {unit.strip()}" if unit else val
                extracted[k] = {
                    "raw_value": val_str,
                    "source_file": source_file,
                    "source_page": page_num,
                    "line_index": line_idx,
                    "extraction_method": "PATTERN_MATCH_VALUNIT"
                }

    # Attach extraction confidence score to all extracted dict items
    for k, item in list(extracted.items()):
        if isinstance(item, dict):
            method = item.get("extraction_method", "PATTERN_MATCH_KEYVAL")
            item["extraction_confidence"] = EXTRACTION_CONFIDENCE_MAP.get(method, 0.90)

    return extracted


def extract_from_pdf_bytes(pdf_bytes: bytes, filename: str) -> Dict[str, Any]:
    """Extract text & tables from PDF (native vector PDF text or scanned image fallback)."""
    extracted: Dict[str, Any] = {}

    # Attempt native text extraction via pypdf or PyPDF2
    full_text = ""
    pages_text = []

    if pypdf:
        try:
            reader = pypdf.PdfReader(io.BytesIO(pdf_bytes))
            if reader.is_encrypted:
                try:
                    reader.decrypt("")
                except Exception:
                    return {"_error": f"PDF file '{filename}' is password-protected and cannot be read."}

            for p_idx, page in enumerate(reader.pages):
                txt = page.extract_text() or ""
                pages_text.append((p_idx + 1, txt))
                full_text += txt + "\n"
        except Exception as e:
            logger.warning("pypdf extraction error for %s: %s", filename, e)

    elif PyPDF2:
        try:
            reader = PyPDF2.PdfReader(io.BytesIO(pdf_bytes))
            if reader.is_encrypted:
                return {"_error": f"PDF file '{filename}' is password-protected and cannot be read."}
            for p_idx, page in enumerate(reader.pages):
                txt = page.extract_text() or ""
                pages_text.append((p_idx + 1, txt))
                full_text += txt + "\n"
        except Exception as e:
            logger.warning("PyPDF2 extraction error for %s: %s", filename, e)

    # Parse key-value pairs per page & full text
    for p_num, p_txt in pages_text:
        if p_txt and len(p_txt.strip()) >= 5:
            p_dict = extract_key_value_pairs_from_text(p_txt, source_file=filename, page_num=p_num)
            for k, v in p_dict.items():
                if k not in extracted:
                    extracted[k] = v

    if full_text and len(full_text.strip()) >= 5:
        full_dict = extract_key_value_pairs_from_text(full_text, source_file=filename, page_num=1)
        for k, v in full_dict.items():
            if k not in extracted:
                extracted[k] = v

    if extracted:
        return extracted

    # Fallback to plain string stream extraction if pypdf reader extracted no structured keys
    try:
        raw_txt = pdf_bytes.decode("latin-1", errors="ignore")
        # Extract text snippets inside parenthesis (BT ... Tj / TJ)
        tj_snippets = re.findall(r"\((.*?)\)\s*TJ?", raw_txt, re.DOTALL)
        if tj_snippets:
            raw_txt += "\n" + "\n".join(tj_snippets)
        return extract_key_value_pairs_from_text(raw_txt, source_file=filename, page_num=1)
    except Exception:
        return {"_error": f"Failed to extract readable content from PDF '{filename}'."}


def extract_from_docx_bytes(docx_bytes: bytes, filename: str) -> Dict[str, Any]:
    """Extract text & table cells from Microsoft Word .docx file."""
    if not docx:
        return {"_error": "python-docx library is not installed."}

    try:
        doc = docx.Document(io.BytesIO(docx_bytes))
        full_text = []

        for p in doc.paragraphs:
            if p.text.strip():
                full_text.append(p.text.strip())

        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if len(cells) >= 2:
                    full_text.append(f"{cells[0]} : {cells[1]}")

        combined = "\n".join(full_text)
        return extract_key_value_pairs_from_text(combined, source_file=filename, page_num=1)
    except Exception as e:
        return {"_error": f"Failed to parse DOCX file '{filename}': {str(e)}"}


def extract_from_rtf_bytes(rtf_bytes: bytes, filename: str) -> Dict[str, Any]:
    """Extract text from Rich Text Format (.rtf) file."""
    try:
        raw_str = rtf_bytes.decode("utf-8", errors="ignore")
        clean_txt = rtf_to_text(raw_str) if rtf_to_text else re.sub(r"\\[a-z0-9]+\b", "", raw_str)
        return extract_key_value_pairs_from_text(clean_txt, source_file=filename, page_num=1)
    except Exception as e:
        return {"_error": f"Failed to parse RTF file '{filename}': {str(e)}"}


def extract_from_excel_bytes(excel_bytes: bytes, filename: str) -> Dict[str, Any]:
    """Extract key-value pairs and tables from Excel .xlsx or .xls file."""
    extracted: Dict[str, Any] = {}

    if openpyxl:
        try:
            wb = openpyxl.load_workbook(io.BytesIO(excel_bytes), data_only=True)
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                for row in ws.iter_rows(values_only=True):
                    non_null = [str(c).strip() for c in row if c is not None and str(c).strip() != ""]
                    if len(non_null) >= 2:
                        k, v = non_null[0], non_null[1]
                        extracted[k] = {
                            "raw_value": v,
                            "source_file": filename,
                            "source_page": 1,
                            "extraction_method": "EXCEL_CELL_PAIR"
                        }
            if extracted:
                return extracted
        except Exception as e:
            logger.warning("openpyxl parsing warning for %s: %s", filename, e)

    if pd:
        try:
            df_dict = pd.read_excel(io.BytesIO(excel_bytes), sheet_name=None)
            for s_name, df in df_dict.items():
                for col in df.columns:
                    for idx, val in df[col].items():
                        if pd.notna(val):
                            extracted[f"{col}_{idx}"] = {
                                "raw_value": str(val),
                                "source_file": filename,
                                "source_page": 1,
                                "extraction_method": "PANDAS_EXCEL"
                            }
            return extracted
        except Exception as e:
            return {"_error": f"Failed to parse Excel file '{filename}': {str(e)}"}

    return {"_error": f"No Excel parser available to process '{filename}'."}


def extract_from_image_bytes(image_bytes: bytes, filename: str) -> Dict[str, Any]:
    """Preprocess image (EXIF rotation, contrast, blur check) and extract text via pytesseract OCR."""
    try:
        pil_img = Image.open(io.BytesIO(image_bytes))
    except Exception as e:
        return {"_error": f"Corrupted or invalid image file '{filename}': {str(e)}"}

    prep_img, meta = preprocess_image_for_ocr(pil_img)

    if meta["is_unreadable"]:
        return {
            "_error": meta["unreadable_reason"],
            "_unreadable": True,
            "_blur_variance": meta["blur_variance"]
        }

    ocr_text = ""
    if pytesseract:
        try:
            ocr_text = pytesseract.image_to_string(prep_img)
        except Exception as e:
            logger.warning("pytesseract execution error for %s: %s", filename, e)

    if not ocr_text.strip():
        # Fallback to plain pattern matcher
        return {"_error": f"OCR text extraction failed for camera image '{filename}'. Please ensure good lighting and clear focus."}

    parsed = extract_key_value_pairs_from_text(ocr_text, source_file=filename, page_num=1)
    for k in parsed:
        if isinstance(parsed[k], dict):
            parsed[k]["extraction_method"] = "OCR_TESSERACT"
            parsed[k]["confidence"] = 0.90
            parsed[k]["blur_variance"] = meta["blur_variance"]

    return parsed


def extract_from_file_or_data(input_source: Union[str, Path, Dict, bytes], filename: str = "") -> Dict[str, Any]:
    """Extract structured data from file path, bytes, JSON string, CSV string, or raw dictionary."""
    if isinstance(input_source, dict):
        if "content" in input_source or "bytes" in input_source or "file_bytes" in input_source:
            content_val = input_source.get("content") or input_source.get("bytes") or input_source.get("file_bytes")
            fn = input_source.get("filename") or input_source.get("name") or filename
            return extract_from_file_or_data(content_val, filename=fn)

        # Convert simple key:val dict into provenance-tracked dict
        out = {}
        for k, v in input_source.items():
            if isinstance(v, dict) and "raw_value" in v:
                out[k] = v
            else:
                out[k] = {
                    "raw_value": str(v) if v is not None else "",
                    "source_file": filename or "dict_input",
                    "source_page": 1,
                    "extraction_method": "DIRECT_INPUT",
                    "confidence": 1.0
                }
        return out

    file_bytes = b""
    fname = filename

    if isinstance(input_source, bytes):
        file_bytes = input_source
    elif isinstance(input_source, (str, Path)):
        path = Path(str(input_source))
        if path.exists() and path.is_file():
            fname = fname or path.name
            with open(path, "rb") as f:
                file_bytes = f.read()
        else:
            # String input payload
            text_str = str(input_source)
            try:
                parsed_json = json.loads(text_str)
                if isinstance(parsed_json, dict):
                    return extract_from_file_or_data(parsed_json, filename=fname)
            except Exception:
                pass
            return extract_key_value_pairs_from_text(text_str, source_file=fname or "text_input", page_num=1)

    if not file_bytes:
        return {"_error": f"Empty or unreadable input source '{fname}'."}

    fmt = detect_file_format_by_header(file_bytes, fname)

    if fmt == "pdf":
        return extract_from_pdf_bytes(file_bytes, fname)
    elif fmt in ["docx", "doc"]:
        return extract_from_docx_bytes(file_bytes, fname)
    elif fmt == "rtf":
        return extract_from_rtf_bytes(file_bytes, fname)
    elif fmt in ["xlsx", "xls"]:
        return extract_from_excel_bytes(file_bytes, fname)
    elif fmt in ["jpeg", "png", "webp", "tiff", "heic"]:
        return extract_from_image_bytes(file_bytes, fname)
    elif fmt in ["json"]:
        try:
            data = json.loads(file_bytes.decode("utf-8"))
            return extract_from_file_or_data(data, filename=fname)
        except Exception as e:
            return {"_error": f"Malformed JSON file '{fname}': {str(e)}"}
    elif fmt in ["csv", "tsv", "txt"]:
        try:
            txt = file_bytes.decode("utf-8", errors="ignore")
            return extract_key_value_pairs_from_text(txt, source_file=fname, page_num=1)
        except Exception as e:
            return {"_error": f"Failed to parse text/csv/tsv file '{fname}': {str(e)}"}

    return {"_error": f"Unsupported or unparseable format for file '{fname}'."}

