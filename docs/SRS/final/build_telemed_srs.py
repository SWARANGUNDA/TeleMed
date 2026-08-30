import os
import sys
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

# Import helper styling module
sys.path.insert(0, r"C:\Users\swara\.gemini\antigravity-ide\brain\255675c7-dd5f-49ee-91ad-966a31752d88\scratch")
from docx_helpers import (
    DIAGRAM_DIR, DOCX_OUTPUT_PATH, HEX_PRIMARY, HEX_SECONDARY, HEX_DARK,
    HEX_LIGHT_BG, HEX_HEADER_BG, HEX_BORDER, COLOR_PRIMARY, COLOR_SECONDARY,
    COLOR_DARK, COLOR_MUTED, set_cell_background, set_cell_margins, set_table_borders,
    add_header_footer, format_paragraph, add_heading_1, add_heading_2, add_heading_3,
    add_body_p, add_bullet_p, add_callout_box, style_table, add_figure
)

def create_srs_document():
    print("Initializing TeleMed SRS Document Creation...")
    doc = Document()

    # Set page margins to 1 inch
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Configure Header and Footer
    add_header_footer(doc)

    # -------------------------------------------------------------------------
    # COVER PAGE
    # -------------------------------------------------------------------------
    cover_p1 = doc.add_paragraph()
    cover_p1.paragraph_format.space_before = Pt(36)
    cover_p1.paragraph_format.space_after = Pt(12)
    cover_p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = cover_p1.add_run("SOFTWARE REQUIREMENTS SPECIFICATION")
    run_title.font.name = "Calibri"
    run_title.font.size = Pt(24)
    run_title.font.bold = True
    run_title.font.color.rgb = COLOR_PRIMARY

    cover_p2 = doc.add_paragraph()
    cover_p2.paragraph_format.space_before = Pt(0)
    cover_p2.paragraph_format.space_after = Pt(24)
    cover_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = cover_p2.add_run("FOR\nGENERATIVE AI ASSISTED TELEMEDICINE PLATFORM FOR PERSONALIZED GUT MICROBIOME ANALYSIS")
    run_sub.font.name = "Calibri"
    run_sub.font.size = Pt(14)
    run_sub.font.bold = True
    run_sub.font.color.rgb = COLOR_SECONDARY

    cover_p3 = doc.add_paragraph()
    cover_p3.paragraph_format.space_before = Pt(12)
    cover_p3.paragraph_format.space_after = Pt(36)
    cover_p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sys = cover_p3.add_run("SYSTEM NAME: TeleMed")
    run_sys.font.name = "Calibri"
    run_sys.font.size = Pt(13)
    run_sys.font.bold = True
    run_sys.font.color.rgb = COLOR_PRIMARY

    # Metadata Box
    add_callout_box(doc, 
        "DOCUMENT IDENTIFIER: SRS-TELEMED-2026-01\n"
        "ACADEMIC DEGREE: Bachelor of Technology (B.Tech) in Computer Science & Engineering\n"
        "CURRICULUM STAGE: III Year / I Semester Major Project Requirement\n"
        "DEPARTMENT: Department of Computer Science & Engineering\n"
        "DATE OF ISSUE: Academic Year 2025–2026\n"
        "SYSTEM STATUS: Production-Grade Research & Decision-Support Prototype",
        "PROJECT SPECIFICATION METADATA"
    )

    doc.add_page_break()

    # -------------------------------------------------------------------------
    # 1. ABSTRACT
    # -------------------------------------------------------------------------
    add_heading_1(doc, "1. ABSTRACT")
    add_body_p(doc, 
        "TeleMed is an advanced AI-assisted telemedicine decision-support platform designed to address the fragmented nature "
        "of contemporary digital health systems. Traditional telemedicine frameworks rely almost exclusively on episodic clinical "
        "laboratory data, failing to synthesize continuous physiological telemetry from wearable sensors and metabolic multi-omics "
        "data from the human gut microbiome. TeleMed resolves this gap by providing a unified, multi-modal diagnostic engine that "
        "integrates three distinct health modalities: (1) Clinical Laboratory & Biomarker Data (18 features), (2) Continuous Wearable "
        "& Continuous Glucose Monitoring (CGM) Telemetry (15 features), and (3) 16S rRNA / Metagenomic Gut Microbiome Profiles "
        "(40 bacterial taxa and 9 derived ecological/functional indices)."
    )
    add_body_p(doc, 
        "The architecture incorporates three specialized modality-specific Machine Learning expert models—a CatBoost Clinical Expert, "
        "an XGBoost Wearable/CGM Expert, and a CatBoost Gut Microbiome Expert. These experts execute modality-level feature processing "
        "and pass their risk estimates into an adaptive Stacking/Fusion Meta-Learner. The fusion engine dynamically evaluates available "
        "patient modalities across seven complete and partial input combinations, producing multi-label risk probability scores for five "
        "interrelated cardiometabolic and liver conditions: Type 2 Diabetes, Prediabetes, High Adiposity Risk, Metabolic Syndrome, and "
        "Non-Alcoholic Fatty Liver Disease (NAFLD)."
    )
    add_body_p(doc, 
        "To ensure transparency and clinical trust, TeleMed integrates TreeSHAP (SHapley Additive exPlanations) to compute exact game-theoretic "
        "feature attributions for every prediction. To bridge the gap between machine predictions and clinical utility, the system combines "
        "a Medical Retrieval-Augmented Generation (RAG) vector engine (utilizing FAISS embeddings of peer-reviewed clinical guidelines) with "
        "a Large Language Model (LLM) to generate personalized, evidence-grounded health reports for patients and reviewing physicians."
    )
    add_callout_box(doc, 
        "CRITICAL DISCLAIMER: TeleMed is designed strictly as an AI-assisted clinical decision-support research prototype. "
        "It does NOT perform autonomous medical diagnosis, prescribe pharmaceutical treatments, or replace qualified healthcare professionals. "
        "All predictions and AI-generated insights represent probabilistic risk estimates and must be interpreted by a certified physician.",
        "CLINICAL DECISION SUPPORT NOTICE"
    )

    # -------------------------------------------------------------------------
    # 2. INTRODUCTION
    # -------------------------------------------------------------------------
    add_heading_1(doc, "2. INTRODUCTION")
    
    add_heading_2(doc, "2.1 Purpose")
    add_body_p(doc, 
        "This Software Requirements Specification (SRS) document provides a complete, rigorous description of the functional, non-functional, "
        "architectural, data, and security requirements for the TeleMed platform. It serves as the authoritative baseline for software engineers, "
        "machine learning architects, quality assurance testers, academic faculty guides, and project stakeholders during system implementation, "
        "testing, and evaluation."
    )

    add_heading_2(doc, "2.2 Problem Statement")
    add_body_p(doc, 
        "Cardiometabolic disorders—including Type 2 Diabetes, Metabolic Syndrome, and NAFLD—represent a global health crisis characterized by "
        "complex multifactorial etiology. Existing telemedicine platforms suffer from significant diagnostic silos:"
    )
    add_bullet_p(doc, "Clinical laboratory panels provide vital snapshot biomarkers (e.g., HbA1c, Fasting Glucose, Lipid Panels) but lack temporal granularity.", "Episodic Snapshot Limitation: ")
    add_bullet_p(doc, "Consumer wearables and CGMs track continuous glycemic variability and autonomic stress but lack biological context regarding organ function or gut dysbiosis.", "Sensor Telemetry Isolation: ")
    add_bullet_p(doc, "Gut microbiome sequencing offers deep insights into systemic inflammation and SCFA production but is rarely integrated into digital clinical workflows.", "Multi-Omics Omission: ")
    add_body_p(doc, 
        "Combining these heterogeneous health modalities into a unified machine learning architecture allows for early risk detection, "
        "personalized lifestyle interventions, and evidence-grounded clinical decision support."
    )

    add_heading_2(doc, "2.3 Objectives")
    add_body_p(doc, "The primary objectives of the TeleMed platform include:")
    objectives = [
        "Architect and implement a secure, scalable web platform for multi-modal health data ingestion.",
        "Develop modality-specific ML expert models for Clinical Data (CatBoost), Wearable Data (XGBoost), and Microbiome Data (CatBoost).",
        "Implement an adaptive 7-pathway Stacking Fusion Meta-Learner supporting complete and partial modality inputs.",
        "Predict multi-label risk probabilities across 5 target conditions (Type 2 Diabetes, Prediabetes, Adiposity, Metabolic Syndrome, NAFLD).",
        "Incorporate TreeSHAP game-theoretic explainability to output precise global and local feature importance attributions.",
        "Integrate a Medical RAG vector search engine (FAISS) grounded in peer-reviewed clinical guidelines.",
        "Deploy Generative AI (LLM) report synthesis with automated safety validation guardrails.",
        "Provide dedicated Patient, Doctor, and Administrator portals with Role-Based Access Control (RBAC).",
        "Generate standardized, downloadable PDF clinical summary reports for patient records.",
        "Maintain immutable assessment history and audit logging for operational governance.",
        "Establish containerized microservices deployment using Docker Compose and PostgreSQL 17.",
        "Ensure full academic rigor and software engineering excellence for III/I B.Tech evaluation."
    ]
    for obj in objectives:
        add_bullet_p(doc, obj)

    add_heading_2(doc, "2.4 Scope")
    add_body_p(doc, "The scope of the TeleMed system is demarcated into clear operational boundaries:")
    add_body_p(doc, "User authentication (JWT/Session), profile management, file upload parsing (PDF/CSV/JSON), unit normalization, physiological bounds validation, 3 ML expert models, 7-pathway fusion, TreeSHAP explanations, FAISS RAG evidence retrieval, LLM report generation, PDF rendering, doctor review workflows, audit logging, and Docker containerization.", "In-Scope Boundaries: ")
    add_body_p(doc, "Autonomous clinical prescription, independent medical diagnosis without doctor oversight, direct hardware Bluetooth syncing, prospective human clinical trial validation, and emergency life-support triage.", "Out-of-Scope Boundaries: ")

    # -------------------------------------------------------------------------
    # 3. FUNCTIONAL REQUIREMENTS
    # -------------------------------------------------------------------------
    add_heading_1(doc, "3. FUNCTIONAL REQUIREMENTS")
    add_body_p(doc, "The functional requirements define the explicit capabilities and behaviors that TeleMed must execute:")

    fr_headers = ["Req ID", "Requirement Name", "Description", "Expected System Behavior"]
    fr_data = [
        ["FR-01", "User Registration", "Allow Patients and Doctors to register accounts.", "Creates user entity, hashes password, initializes role profile."],
        ["FR-02", "User Authentication", "Authenticate users via email and password.", "Validates credentials, issues JWT access/refresh tokens and HTTP cookies."],
        ["FR-03", "Role-Based Access Control", "Enforce RBAC boundaries (Patient, Doctor, Admin).", "Blocks unauthorized endpoint access based on role permissions."],
        ["FR-04", "Patient Profile", "Maintain demographic and anthropometric records.", "Stores and calculates derived metrics such as BMI from height/weight."],
        ["FR-05", "Clinical Input", "Ingest 18 clinical laboratory biomarkers.", "Validates numeric inputs against physiological boundaries."],
        ["FR-06", "Wearable/CGM Input", "Ingest 15 wearable sensor & CGM features.", "Parses CSV telemetry files and populates wearable feature tensor."],
        ["FR-07", "Microbiome Input", "Ingest 40 bacterial taxa & 9 derived indices.", "Validates relative abundance sum and computes ecological diversity."],
        ["FR-08", "File Upload Parsing", "Upload medical reports in PDF, CSV, TXT, PNG.", "Extracts structured key-value pairs via IMDIE parser engine."],
        ["FR-09", "Data Validation", "Check physiological bounds and alias mapping.", "Flags anomalies and resolves canonical feature aliases automatically."],
        ["FR-10", "Unit Normalization", "Standardize clinical and wearable units.", "Converts mg/dL to mmol/L or standard units prior to inference."],
        ["FR-11", "Modality Detection", "Detect available input modalities dynamically.", "Identifies active modalities and routes tensor to corresponding pathway."],
        ["FR-12", "Clinical Expert ML", "Execute CatBoost Clinical Expert model.", "Computes clinical-specific risk probabilities."],
        ["FR-13", "Wearable Expert ML", "Execute XGBoost Wearable/CGM Expert model.", "Computes wearable-specific risk probabilities."],
        ["FR-14", "Gut Expert ML", "Execute CatBoost Gut Microbiome Expert model.", "Computes microbiome-specific risk probabilities."],
        ["FR-15", "Multimodal Fusion", "Execute Stacking Meta-Learner.", "Combines expert predictions across 1 of 7 active pathways."],
        ["FR-16", "Five-Disease Risk", "Generate multi-label risk probability scores.", "Outputs 5 independent risk scores (0.0 to 1.0) for target conditions."],
        ["FR-17", "Risk Categorization", "Classify risk scores into LOW, MODERATE, HIGH.", "Applies clinical threshold boundaries for dashboard presentation."],
        ["FR-18", "TreeSHAP Attributions", "Calculate local & global Shapley values.", "Displays top positive and negative feature contributions per prediction."],
        ["FR-19", "Prediction Persistence", "Store assessment snapshots in PostgreSQL.", "Persists feature inputs, predictions, and SHAP attributions."],
        ["FR-20", "Medical RAG Retrieval", "Query FAISS vector store with SHAP features.", "Retrieves top-k relevant clinical guideline excerpts."],
        ["FR-21", "Generative AI Report", "Synthesize evidence-grounded health report.", "LLM generates patient summary and doctor clinical briefing."],
        ["FR-22", "Report Guardrails", "Validate LLM output against safety rules.", "Rejects reports containing ungrounded claims or diagnostic assertions."],
        ["FR-23", "PDF Generation", "Render downloadable PDF assessment report.", "Generates styled PDF containing metrics, SHAP charts, and report."],
        ["FR-24", "Assessment History", "Display historical assessment timeline.", "Allows patients and doctors to view prior assessment progression."],
        ["FR-25", "Doctor Portal", "Enable doctor review and verification.", "Doctors inspect assigned patient assessments and append review notes."],
        ["FR-26", "Admin Governance", "Manage users, doctor verification, system health.", "Provides admin dashboard for system monitoring and account governance."],
        ["FR-27", "Error Handling", "Handle missing modalities & server errors.", "Displays descriptive user alerts and maintains operational stability."],
        ["FR-28", "Audit Logging", "Record security and data access audit events.", "Logs login attempts, data access, and clinical status transitions."]
    ]
    t_fr = doc.add_table(rows=1, cols=4)
    style_table(t_fr, [1.0, 1.6, 2.2, 1.7], fr_headers, fr_data)

    # -------------------------------------------------------------------------
    # 4. NON-FUNCTIONAL REQUIREMENTS
    # -------------------------------------------------------------------------
    add_heading_1(doc, "4. NON-FUNCTIONAL REQUIREMENTS")
    add_body_p(doc, "Non-functional requirements specify quality attributes, performance benchmarks, security standards, and system constraints:")

    nfr_headers = ["ID", "Category", "Requirement Specification", "Target Metric / Standard"]
    nfr_data = [
        ["NFR-01", "Performance", "End-to-end ML inference response time.", "< 1.5 seconds for multi-modal inference."],
        ["NFR-02", "Performance", "RAG vector search and LLM report synthesis.", "< 4.0 seconds for complete report generation."],
        ["NFR-03", "Availability", "System operational uptime.", "99.5% availability during operational hours."],
        ["NFR-04", "Scalability", "Stateless API server design with connection pooling.", "Supports 100+ concurrent user sessions."],
        ["NFR-05", "Security", "Data in transit and at rest encryption.", "TLS 1.3 for in-transit; AES-256 for database storage."],
        ["NFR-06", "Security", "Password storage hashing.", "PBKDF2 / SHA-256 with unique 16-byte random salt."],
        ["NFR-07", "Privacy", "Patient clinical data isolation.", "Strict tenant boundaries enforced at API & DB query level."],
        ["NFR-08", "Usability", "Responsive, high-aesthetic web interface.", "Fully functional across desktop, tablet, and mobile browsers."],
        ["NFR-09", "Maintainability", "Modular code structure with separation of concerns.", "Clean REST architecture with documented FastAPI endpoints."],
        ["NFR-10", "Extensibility", "Pluggable ML expert model architecture.", "Allows new modality experts to be added without breaking pipeline."],
        ["NFR-11", "Explainability", "Interpretable prediction attributions.", "100% of predictions accompanied by TreeSHAP feature scores."],
        ["NFR-12", "Safety", "AI non-autonomy and safety disclaimers.", "Mandatory medical disclaimer on all reports and UI views."],
        ["NFR-13", "Data Integrity", "Strict schema validation and bounds checking.", "Rejects invalid physiological inputs prior to ML execution."],
        ["NFR-14", "Auditability", "Comprehensive audit logging of data operations.", "All clinical data access logged with timestamp and user ID."],
        ["NFR-15", "Reproducibility", "Containerized software deployment environment.", "Docker Compose deployment with pinned library dependencies."]
    ]
    t_nfr = doc.add_table(rows=1, cols=4)
    style_table(t_nfr, [0.8, 1.3, 2.8, 1.6], nfr_headers, nfr_data)

    # -------------------------------------------------------------------------
    # 5. IDENTIFICATION OF USERS
    # -------------------------------------------------------------------------
    add_heading_1(doc, "5. IDENTIFICATION OF USERS")
    add_body_p(doc, "The TeleMed platform defines three primary human user roles and one external system interface:")

    add_heading_2(doc, "5.1 Patient User")
    add_body_p(doc, "Primary consumer of the platform seeking personalized health assessments. Responsibilities include registering, managing personal demographics, submitting clinical lab values, uploading wearable/CGM files, uploading microbiome datasets, initiating assessments, viewing risk dashboard, inspecting TreeSHAP explanations, and downloading PDF clinical summaries.")

    add_heading_2(doc, "5.2 Doctor User (Clinical Reviewer)")
    add_body_p(doc, "Licensed medical practitioner responsible for clinical oversight. Responsibilities include logging in via verified credentials, accessing assigned patient records, inspecting multi-modal risk profiles and SHAP attributions, evaluating AI-generated clinical briefings, adding physician notes, and guiding patient care plans. The doctor retains full clinical responsibility.")

    add_heading_2(doc, "5.3 Administrator User")
    add_body_p(doc, "System operations and governance user. Responsibilities include managing user accounts, verifying doctor registration credentials, monitoring system health and latency, inspecting audit logs, and maintaining platform security.")

    add_heading_2(doc, "5.4 External Medical Knowledge Base Interface")
    add_body_p(doc, "External vector database storing peer-reviewed clinical practice guidelines (ADA, EASD, AASLD). Interfaced via FAISS vector similarity search during RAG evidence retrieval.")

    # -------------------------------------------------------------------------
    # 6. MODULES IN THE APPLICATION
    # -------------------------------------------------------------------------
    add_heading_1(doc, "6. MODULES IN THE APPLICATION")
    add_body_p(doc, "TeleMed is structured into 16 cohesive software modules:")

    modules_data = [
        ("1. Authentication & RBAC Module", "Manages user registration, password hashing, JWT token issuance, session tracking, and role enforcement.", "Interacts with User DB, API Gateway."),
        ("2. Patient Management Module", "Handles patient demographic profiles, anthropometric calculations (BMI), and personal health preferences.", "Interacts with Patient DB, Intake Module."),
        ("3. Doctor Management Module", "Manages doctor registration, medical license verification workflows, and assigned patient lists.", "Interacts with Doctor DB, Review Portal."),
        ("4. Multimodal Data Intake Module", "Parses uploaded PDF, CSV, TXT, and PNG files to extract structured key-value feature pairs.", "Interacts with Upload Parser, Validation Module."),
        ("5. Validation & Preprocessing Module", "Validates physiological bounds, converts measurement units, and normalizes feature tensors.", "Interacts with Intake Module, ML Experts."),
        ("6. Clinical Expert Module", "Executes CatBoost model trained on 18 clinical laboratory biomarkers.", "Interacts with Validation Module, Fusion Engine."),
        ("7. Wearable/CGM Expert Module", "Executes XGBoost model trained on 15 wearable sensor and CGM features.", "Interacts with Validation Module, Fusion Engine."),
        ("8. Gut Microbiome Expert Module", "Executes CatBoost model trained on 40 bacterial taxa and 9 ecological/functional indices.", "Interacts with Validation Module, Fusion Engine."),
        ("9. Multimodal Fusion Module", "Implements Stacking Meta-Learner dynamically selecting 1 of 7 modality pathways.", "Interacts with Expert Models, XAI Module."),
        ("10. Explainable AI (TreeSHAP) Module", "Calculates exact Shapley value attributions for global and local feature explanations.", "Interacts with Fusion Engine, Database."),
        ("11. Medical RAG Retrieval Module", "Queries FAISS vector index with top SHAP features to retrieve clinical guideline excerpts.", "Interacts with XAI Module, LLM Module."),
        ("12. Generative AI Synthesis Module", "Generates structured, evidence-grounded health reports using LLM prompt templates.", "Interacts with RAG Module, Safety Guardrails."),
        ("13. Report & PDF Engine Module", "Formats clinical assessment outputs and renders publication-quality PDF documents.", "Interacts with GenAI Module, Patient/Doctor UI."),
        ("14. Assessment History Module", "Persists and queries historical assessment snapshots for longitudinal trend analysis.", "Interacts with Database, Dashboard UI."),
        ("15. Database & Persistence Module", "Provides SQLAlchemy ORM layer for PostgreSQL 17 transaction and entity management.", "Interacts with All Backend Services."),
        ("16. Administration & Monitoring Module", "Provides operational dashboards, audit logging, Prometheus metrics, and system status.", "Interacts with System Logs, Admin UI.")
    ]
    for name, resp, inter in modules_data:
        add_body_p(doc, f"{resp} ({inter})", bold_prefix=f"{name}: ")

    # -------------------------------------------------------------------------
    # 7. DATA REQUIREMENTS
    # -------------------------------------------------------------------------
    add_heading_1(doc, "7. DATA REQUIREMENTS & FEATURE SPECIFICATIONS")
    
    add_heading_2(doc, "7.1 Dataset Overview")
    add_body_p(doc, 
        "TeleMed processes 82 total modality input features spanning Clinical (18), Wearable/CGM (15), and Gut Microbiome (49 model input tensor features). "
        "The system predicts 5 multi-label target conditions. Patient_ID is treated purely as metadata and is never passed to ML models."
    )

    add_heading_2(doc, "7.2 Clinical Modality Features (18 Features)")
    clin_headers = ["#", "Feature Name", "Category", "Description", "Data Type / Unit"]
    clin_data = [
        ["1", "Age", "Demographic", "Patient age in years", "Integer (years)"],
        ["2", "Gender", "Demographic", "Biological sex at birth", "Categorical (Male/Female)"],
        ["3", "Height", "Anthropometric", "Body height", "Float (cm)"],
        ["4", "Weight", "Anthropometric", "Body weight", "Float (kg)"],
        ["5", "BMI", "Anthropometric", "Body Mass Index", "Float (kg/m²)"],
        ["6", "Waist_Circumference", "Anthropometric", "Waist circumference", "Float (cm)"],
        ["7", "Systolic_BP", "Vascular", "Systolic blood pressure", "Float (mmHg)"],
        ["8", "Diastolic_BP", "Vascular", "Diastolic blood pressure", "Float (mmHg)"],
        ["9", "Fasting_Blood_Glucose", "Metabolic", "Fasting plasma glucose", "Float (mg/dL)"],
        ["10", "HbA1c", "Metabolic", "Glycated hemoglobin", "Float (%)"],
        ["11", "Triglycerides", "Lipid Panel", "Serum triglycerides", "Float (mg/dL)"],
        ["12", "HDL", "Lipid Panel", "High-density lipoprotein", "Float (mg/dL)"],
        ["13", "LDL", "Lipid Panel", "Low-density lipoprotein", "Float (mg/dL)"],
        ["14", "ALT", "Hepatic", "Alanine aminotransferase", "Float (U/L)"],
        ["15", "AST", "Hepatic", "Aspartate aminotransferase", "Float (U/L)"],
        ["16", "Family_History_Diabetes", "Genetic Risk", "Family history of T2D", "Binary (0/1)"],
        ["17", "Family_History_Hypertension", "Genetic Risk", "Family history of HTN", "Binary (0/1)"],
        ["18", "Family_History_CVD", "Genetic Risk", "Family history of CVD", "Binary (0/1)"]
    ]
    t_clin = doc.add_table(rows=1, cols=5)
    style_table(t_clin, [0.4, 1.8, 1.2, 2.1, 1.2], clin_headers, clin_data)

    add_heading_2(doc, "7.3 Wearable & CGM Modality Features (15 Features)")
    wear_headers = ["#", "Feature Name", "Source", "Description", "Data Type / Unit"]
    wear_data = [
        ["1", "Average_Daily_Steps", "Wearable Sensor", "Mean daily step counts over tracking period", "Integer (steps)"],
        ["2", "Active_Minutes", "Wearable Sensor", "Daily moderate-to-vigorous physical activity", "Float (minutes)"],
        ["3", "Sedentary_Time_Minutes", "Wearable Sensor", "Daily inactive duration", "Float (minutes)"],
        ["4", "Resting_Heart_Rate", "Wearable Sensor", "Baseline resting heart rate", "Float (bpm)"],
        ["5", "Heart_Rate_Variability_RMSSD", "Wearable Sensor", "Root mean square of successive RR diffs", "Float (ms)"],
        ["6", "Sleep_Duration_Hours", "Wearable Sensor", "Total sleep duration per 24h cycle", "Float (hours)"],
        ["7", "Sleep_Efficiency_Score", "Wearable Sensor", "Ratio of time asleep to time in bed", "Float (0.0–1.0)"],
        ["8", "Autonomic_Stress_Score", "Wearable Sensor", "Derived autonomic stress burden", "Float (0–100 index)"],
        ["9", "Activity_Energy_Expenditure", "Wearable Sensor", "Caloric expenditure from activity", "Float (kcal)"],
        ["10", "Exercise_Frequency_Days", "Wearable Sensor", "Days per week with structured exercise", "Integer (days)"],
        ["11", "CGM_Average_Glucose", "Continuous Glucose", "24-hour mean interstitial glucose", "Float (mg/dL)"],
        ["12", "CGM_Glucose_CV", "Continuous Glucose", "Coefficient of variation in glucose", "Float (%)"],
        ["13", "CGM_Time_In_Range", "Continuous Glucose", "Time spent between 70–180 mg/dL", "Float (%)"],
        ["14", "CGM_Time_Above_Range", "Continuous Glucose", "Time spent > 180 mg/dL", "Float (%)"],
        ["15", "CGM_Time_Below_Range", "Continuous Glucose", "Time spent < 70 mg/dL", "Float (%)"]
    ]
    t_wear = doc.add_table(rows=1, cols=5)
    style_table(t_wear, [0.4, 2.2, 1.2, 1.7, 1.2], wear_headers, wear_data)

    add_heading_2(doc, "7.4 Gut Microbiome Modality Features (49 Model Tensor Inputs / 50 Quantities)")
    add_body_p(doc, 
        "The Gut Microbiome Expert processes 40 bacterial taxa, 1 residual compositional component (Other_Taxa), and 9 derived "
        "ecological and functional indices. The model tensor input contains 49 features."
    )
    
    add_callout_box(doc, 
        "SCHEMA VERIFICATION NOTE: The biological dataset defines 40 named taxa, 1 Other_Taxa component, and 9 derived indices "
        "(totaling 50 biological quantities). In the ML tensor representation, the 49 model input features comprise the 40 taxa, "
        "Other_Taxa, and 8 primary ecological indices (with 1 collinear index reserved for downstream verification). "
        "All measurements represent relative abundances (summing to 100%) or normalized index metrics.",
        "DATA SCHEMA CONSISTENCY NOTE"
    )

    gut_headers = ["#", "Feature Name", "Type / Category", "Biological & Functional Description"]
    gut_data = [
        ["1", "Akkermansia_muciniphila", "Bacterial Taxon", "Mucin-degrading bacterium associated with gut barrier integrity & insulin sensitivity."],
        ["2", "Faecalibacterium_prausnitzii", "Bacterial Taxon", "Major butyrate producer with potent anti-inflammatory properties."],
        ["3", "Roseburia_intestinalis", "Bacterial Taxon", "Short-chain fatty acid (SCFA) producer involved in metabolic regulation."],
        ["4", "Bifidobacterium_longum", "Bacterial Taxon", "Beneficial probiotic taxon supporting immune homeostasis."],
        ["5", "Bifidobacterium_adolescentis", "Bacterial Taxon", "Associated with carbohydrate fermentation and acetate production."],
        ["6", "Bacteroides_thetaiotaomicron", "Bacterial Taxon", "Complex polysaccharide degrader and mucosal symbiont."],
        ["7", "Bacteroides_vulgatus", "Bacterial Taxon", "Prevalent gut commensal linked to lipid metabolism modulation."],
        ["8", "Bacteroides_fragilis", "Bacterial Taxon", "Immunomodulatory commensal involved in T-cell maturation."],
        ["9", "Bacteroides_uniformis", "Bacterial Taxon", "Associated with fiber digestion and reduced metabolic inflammation."],
        ["10", "Prevotella_copri", "Bacterial Taxon", "Fiber-fermenting taxon associated with dietary carbohydrate responsiveness."],
        ["11", "Ruminococcus_bromii", "Bacterial Taxon", "Primary resistant-starch degrader producing fermentable substrates."],
        ["12", "Ruminococcus_gnavus", "Bacterial Taxon", "Inflammatory-associated taxon reported in metabolic dysbiosis."],
        ["13", "Blautia_wexlerae", "Bacterial Taxon", "Commensal associated with metabolic health and visceral fat regulation."],
        ["14", "Blautia_hansenii", "Bacterial Taxon", "Acetogenic taxon participating in SCFA cross-feeding networks."],
        ["15", "Collinsella_aerofaciens", "Bacterial Taxon", "Correlated with serum cholesterol metabolism and bile acid modification."],
        ["16", "Escherichia_coli", "Bacterial Taxon", "Proteobacterial marker whose expansion indicates gut dysbiosis."],
        ["17", "Klebsiella_pneumoniae", "Bacterial Taxon", "Potential pathobiont associated with endotoxemia and systemic inflammation."],
        ["18", "Coprococcus_eutactus", "Bacterial Taxon", "Butyrate producer associated with positive metabolic outcomes."],
        ["19", "Alistipes_putredinis", "Bacterial Taxon", "Bacteroidetes member involved in protein and bile acid metabolism."],
        ["20", "Alistipes_finegoldii", "Bacterial Taxon", "Associated with gut gut environment stability and lipid processing."],
        ["21", "Subdoligranulum_variable", "Bacterial Taxon", "Butyrate-producing firmicute correlated with insulin sensitivity."],
        ["22", "Enterococcus_faecalis", "Bacterial Taxon", "Opportunistic commensal monitored as a microbial balance indicator."],
        ["23", "Eubacterium_rectale", "Bacterial Taxon", "Major SCFA producer playing a central role in colonic health."],
        ["24", "Eubacterium_hallii", "Bacterial Taxon", "Metabolizes lactate and acetate into butyrate."],
        ["25", "Parabacteroides_distasonis", "Bacterial Taxon", "Associated with succinate production and metabolic regulation."],
        ["26", "Lactobacillus_acidophilus", "Bacterial Taxon", "Lactic acid producer supporting intestinal barrier functionality."],
        ["27", "Lactobacillus_rhamnosus", "Bacterial Taxon", "Probiotic strain associated with reduced intestinal permeability."],
        ["28", "Streptococcus_thermophilus", "Bacterial Taxon", "Dietary/fermentative microbe contributing to lactic acid production."],
        ["29", "Eggerthella_lenta", "Bacterial Taxon", "Actinobacterial strain involved in xenobiotic and cardenolide metabolism."],
        ["30", "Christensenella_minuta", "Bacterial Taxon", "Heritable microbe strongly associated with low BMI and metabolic health."],
        ["31", "Methanobrevibacter_smithii", "Bacterial Taxon", "Dominant methanogenic archaeon influencing caloric extraction efficiency."],
        ["32", "Dialister_invisus", "Bacterial Taxon", "Metabolic commensal monitored in glucose tolerance studies."],
        ["33", "Holdemanella_biformis", "Bacterial Taxon", "Produces long-chain fatty acid derivatives influencing GLP-1 secretion."],
        ["34", "Barnesiella_intestinihominis", "Bacterial Taxon", "Immunomodulatory commensal supporting intestinal clearance."],
        ["35", "Anaerostipes_caccae", "Bacterial Taxon", "Converts lactate into butyrate, contributing to mucosal nourishment."],
        ["36", "Phascolarctobacterium_faecium", "Bacterial Taxon", "Produces propionate from succinate utilization."],
        ["37", "Veillonella_parvula", "Bacterial Taxon", "Lactate-utilizing microbe linked to exercise physiology."],
        ["38", "Fusobacterium_nucleatum", "Bacterial Taxon", "Pro-inflammatory pathobiont associated with mucosal barrier breach."],
        ["39", "Bilophila_wadsworthia", "Bacterial Taxon", "Sulfite-reducing microbe associated with high-fat diets and hydrogen sulfide."],
        ["40", "Sutterella_wadsworthensis", "Bacterial Taxon", "Proteobacterial commensal monitored in inflammatory bowel profiles."],
        ["41", "Other_Taxa", "Residual Component", "Aggregate relative abundance of remaining low-abundance microflora."],
        ["42", "Shannon_Diversity", "Ecological Index", "Alpha diversity index measuring species richness and evenness."],
        ["43", "Simpson_Diversity", "Ecological Index", "Alpha diversity index emphasizing dominant microbial species."],
        ["44", "Observed_Richness", "Ecological Index", "Total count of unique bacterial taxa identified in sample."],
        ["45", "Pielou_Evenness", "Ecological Index", "Measure of how equally species abundances are distributed."],
        ["46", "SCFA_Producer_Index", "Functional Index", "Aggregate functional capacity for total short-chain fatty acid synthesis."],
        ["47", "Butyrate_Producer_Index", "Functional Index", "Derived score reflecting abundance of primary butyrate-producing taxa."],
        ["48", "Barrier_Associated_Index", "Functional Index", "Composite index of mucin-preserving and barrier-supporting species."],
        ["49", "Inflammation_Index", "Functional Index", "Ratio of pro-inflammatory pathobionts to anti-inflammatory commensals."],
        ["50", "Log_F_B_Ratio", "Functional Index", "Log-transformed ratio of Firmicutes to Bacteroidetes phyla."]
    ]
    t_gut = doc.add_table(rows=1, cols=4)
    style_table(t_gut, [0.4, 2.2, 1.4, 2.7], gut_headers, gut_data)

    add_heading_2(doc, "7.5 Target Disease Conditions (5 Multi-Label Targets)")
    target_headers = ["Condition Name", "Clinical Definition & Criteria", "Primary Associated Modalities"]
    target_data = [
        ["Type 2 Diabetes", "HbA1c ≥ 6.5% or Fasting Glucose ≥ 126 mg/dL or CGM Mean > 180 mg/dL.", "Clinical, CGM, Gut Microbiome"],
        ["Prediabetes", "HbA1c 5.7–6.4% or Fasting Glucose 100–125 mg/dL.", "Clinical, CGM, Wearables"],
        ["High Adiposity Risk", "BMI ≥ 27.5 kg/m² or Waist Circumference > 90cm (M) / 80cm (F).", "Clinical, Wearables, Gut Taxa"],
        ["Metabolic Syndrome", "Presence of ≥3 criteria: Elevated BP, Glucose, Triglycerides, Low HDL, Waist.", "Clinical, Wearables, Gut Indices"],
        ["NAFLD Risk", "Elevated ALT/AST ratio, High BMI, Triglycerides, and Dysbiosis markers.", "Clinical (Hepatic), Gut Microbiome"]
    ]
    t_target = doc.add_table(rows=1, cols=3)
    style_table(t_target, [1.8, 2.9, 2.0], target_headers, target_data)

    # -------------------------------------------------------------------------
    # 8. MACHINE LEARNING ARCHITECTURE
    # -------------------------------------------------------------------------
    add_heading_1(doc, "8. MACHINE LEARNING ARCHITECTURE")
    add_body_p(doc, 
        "TeleMed implements a decoupled multi-expert architecture. Instead of concatenating raw features from different domains "
        "into a single hyper-sparse vector, dedicated expert algorithms are trained on domain-specific feature tensors:"
    )
    add_bullet_p(doc, "Gradient boosted decision trees optimized for tabular lab biomarkers and clinical risk factors. Handles missing clinical entries natively via symmetric trees.", "Clinical Expert (CatBoost): ")
    add_bullet_p(doc, "Gradient boosted trees optimized for continuous physiological sensor metrics and CGM glycemic dynamics.", "Wearable/CGM Expert (XGBoost): ")
    add_bullet_p(doc, "Gradient boosted trees optimized for non-linear compositional microbiome abundances and non-Euclidean ecological index spaces.", "Gut Microbiome Expert (CatBoost): ")

    # -------------------------------------------------------------------------
    # 9. MULTIMODAL FUSION
    # -------------------------------------------------------------------------
    add_heading_1(doc, "9. MULTIMODAL FUSION & 7-PATHWAY ENGINE")
    add_body_p(doc, 
        "Patients frequently present with incomplete datasets (e.g., submitting lab results and wearable data without gut microbiome sequencing). "
        "TeleMed solves this using a Stacking Meta-Learner that dynamically routes available data across 7 distinct fusion pathways:"
    )

    pathway_headers = ["Pathway ID", "Active Modalities", "Input Vector to Fusion Engine", "Fallback / Execution Strategy"]
    pathway_data = [
        ["Pathway 1", "Clinical Only", "Clinical Expert Probabilities [5]", "Direct Clinical Expert output scaling."],
        ["Pathway 2", "Wearable/CGM Only", "Wearable Expert Probabilities [5]", "Direct Wearable Expert output scaling."],
        ["Pathway 3", "Gut Microbiome Only", "Gut Expert Probabilities [5]", "Direct Microbiome Expert output scaling."],
        ["Pathway 4", "Clinical + Wearable", "Concatenated [Clinical_Prob, Wearable_Prob]", "Meta-Learner Stacking Model 4."],
        ["Pathway 5", "Clinical + Gut", "Concatenated [Clinical_Prob, Gut_Prob]", "Meta-Learner Stacking Model 5."],
        ["Pathway 6", "Wearable + Gut", "Concatenated [Wearable_Prob, Gut_Prob]", "Meta-Learner Stacking Model 6."],
        ["Pathway 7", "Clinical + Wearable + Gut", "Concatenated [Clin_Prob, Wear_Prob, Gut_Prob]", "Full Multimodal Stacking Meta-Learner 7."]
    ]
    t_path = doc.add_table(rows=1, cols=4)
    style_table(t_path, [1.1, 1.8, 2.3, 1.5], pathway_headers, pathway_data)

    # -------------------------------------------------------------------------
    # 10. EXPLAINABLE AI (XAI)
    # -------------------------------------------------------------------------
    add_heading_1(doc, "10. EXPLAINABLE AI (TreeSHAP)")
    add_body_p(doc, 
        "To eliminate the 'black-box' nature of gradient boosted ensembles, TeleMed integrates TreeSHAP (SHapley Additive exPlanations). "
        "TreeSHAP calculates exact additive feature attributions based on cooperative game theory. For every prediction, the system outputs "
        "local feature impact scores explaining exactly how much each biomarker or bacterial taxon pushed the risk score above or below the baseline population average."
    )

    # -------------------------------------------------------------------------
    # 11. MEDICAL RAG
    # -------------------------------------------------------------------------
    add_heading_1(doc, "11. MEDICAL RAG ENGINE")
    add_body_p(doc, 
        "The Medical Retrieval-Augmented Generation (RAG) module grounds AI reports in peer-reviewed medical literature. "
        "Peer-reviewed clinical guidelines (ADA Standards of Care, EASD, AASLD NAFLD Guidelines) are chunked, embedded using vector models, "
        "and stored in a FAISS vector index. When an assessment completes, top TreeSHAP risk factors query the FAISS index to retrieve "
        "the top-k (k=3) relevant clinical evidence passages."
    )

    # -------------------------------------------------------------------------
    # 12. GENERATIVE AI
    # -------------------------------------------------------------------------
    add_heading_1(doc, "12. GENERATIVE AI REPORT GENERATION")
    add_body_p(doc, 
        "The Generative AI module uses an LLM to synthesize patient-friendly health summaries and physician clinical briefings. "
        "The prompt payload includes: (1) Patient Demographics, (2) Calculated Risk Probabilities, (3) Top 5 TreeSHAP Contributing Features, "
        "and (4) Retrieved Medical RAG Evidence Passages."
    )
    add_callout_box(doc, 
        "SAFETY GUARDRAIL ENFORCEMENT: All generated reports pass through automated regex and structural validators. "
        "Reports containing forbidden phrases (e.g., 'diagnosed with', 'prescribe medication', 'guaranteed cure') or ungrounded claims "
        "are flagged and rejected automatically.",
        "GENERATIVE AI SAFETY GUARDRAILS"
    )

    # -------------------------------------------------------------------------
    # 13. DATABASE REQUIREMENTS
    # -------------------------------------------------------------------------
    add_heading_1(doc, "13. DATABASE & PERSISTENCE REQUIREMENTS")
    add_body_p(doc, "TeleMed utilizes PostgreSQL 17 as its primary relational database engine. Core domain entities include:")
    add_bullet_p(doc, "Stores account credentials, role (PATIENT, DOCTOR, ADMIN), and authentication metadata.", "User Entity: ")
    add_bullet_p(doc, "Stores demographic profiles, height, weight, BMI, and contact information.", "PatientProfile Entity: ")
    add_bullet_p(doc, "Stores medical license numbers, specialization, hospital affiliation, and verification status.", "DoctorProfile Entity: ")
    add_bullet_p(doc, "Records assessment session timestamps, active pathway ID, and overall data quality scores.", "HealthAssessment Entity: ")
    add_bullet_p(doc, "Stores 5 target disease risk probability scores and category labels.", "PredictionResult Entity: ")
    add_bullet_p(doc, "Stores JSON-serialized TreeSHAP local feature attribution maps.", "XAIExplanation Entity: ")
    add_bullet_p(doc, "Stores AI-generated clinical reports, RAG provenance metadata, and PDF storage paths.", "ClinicalReport Entity: ")

    # -------------------------------------------------------------------------
    # 14. SYSTEM ARCHITECTURE
    # -------------------------------------------------------------------------
    add_heading_1(doc, "14. SYSTEM ARCHITECTURE & TECH STACK")
    add_body_p(doc, "The platform is implemented using a modern multi-tier enterprise stack:")
    add_bullet_p(doc, "React 18 + Vite, TailwindCSS / Custom CSS Design Tokens, Lucide Icons.", "Frontend Framework: ")
    add_bullet_p(doc, "FastAPI (Python 3.12), Pydantic v2 validation, Uvicorn ASGI server.", "Backend REST Framework: ")
    add_bullet_p(doc, "PostgreSQL 17 with SQLAlchemy 2.0 ORM and Alembic migrations.", "Database Engine: ")
    add_bullet_p(doc, "CatBoost 1.2+, XGBoost 2.0+, Scikit-Learn 1.4+, TreeSHAP 0.44+.", "Machine Learning Core: ")
    add_bullet_p(doc, "FAISS (Facebook AI Similarity Search), SentenceTransformers, LLM Integration.", "Vector RAG & GenAI: ")
    add_bullet_p(doc, "Docker, Docker Compose, Redis, Prometheus, Grafana.", "Containerization & Ops: ")

    # -------------------------------------------------------------------------
    # 15. UML DIAGRAMS
    # -------------------------------------------------------------------------
    add_heading_1(doc, "15. SYSTEM UML DIAGRAMS")
    add_body_p(doc, "This section presents the foundational Object-Oriented and Structural UML models governing TeleMed:")

    add_heading_2(doc, "15.1 Use Case Diagram")
    add_body_p(doc, "The Use Case Diagram defines the interactions between human actors (Patient, Doctor, Administrator) and the TeleMed system boundary.")
    add_figure(doc, "use_case_diagram.png", "Figure 15.1: TeleMed Platform Use Case Model Diagram")

    add_heading_2(doc, "15.2 Class Diagram")
    add_body_p(doc, "The Class Diagram models the domain entities, ML expert classes, fusion engine, and database relationships.")
    add_figure(doc, "class_diagram.png", "Figure 15.2: TeleMed Domain & Class Architecture Diagram")

    add_heading_2(doc, "15.3 Sequence Diagram")
    add_body_p(doc, "The Sequence Diagram specifies the chronological message flow during a patient health assessment and report generation lifecycle.")
    add_figure(doc, "sequence_diagram.png", "Figure 15.3: Patient Health Assessment & AI Report Generation Sequence Diagram")

    add_heading_2(doc, "15.4 Component Diagram")
    add_body_p(doc, "The Component Diagram details the multi-tier architectural layers from UI Presentation down to ML Core and Database Persistence.")
    add_figure(doc, "component_diagram.png", "Figure 15.4: TeleMed Multi-Tier Component & Layer Architecture Diagram")

    # -------------------------------------------------------------------------
    # 16. DATA FLOW
    # -------------------------------------------------------------------------
    add_heading_1(doc, "16. DATA FLOW ARCHITECTURE")
    add_body_p(doc, 
        "Data flows sequentially through 6 operational stages: (1) Ingestion & Parse -> (2) Bounds Check & Unit Normalization -> "
        "(3) Modality Expert ML Execution -> (4) Stacking Fusion Meta-Inference -> (5) TreeSHAP Attribution & RAG Retrieval -> "
        "(6) GenAI Report Formatting & PostgreSQL Persistence."
    )

    # -------------------------------------------------------------------------
    # 17. SECURITY, PRIVACY AND SAFETY
    # -------------------------------------------------------------------------
    add_heading_1(doc, "17. SECURITY, PRIVACY, AND AI SAFETY")
    add_body_p(doc, "Security measures include JWT bearer tokens, bcrypt/PBKDF2 password hashing, strict RBAC tenant isolation, CORS origin control, SQL injection prevention via SQLAlchemy parameterized queries, and mandatory medical AI safety disclaimers.")

    # -------------------------------------------------------------------------
    # 18. TESTING REQUIREMENTS
    # -------------------------------------------------------------------------
    add_heading_1(doc, "18. TESTING & EVALUATION REQUIREMENTS")
    add_body_p(doc, "Testing encompasses Unit Tests (PyTest), API Integration Tests, Security Penetration Scenarios, and ML Evaluation Metrics:")
    add_bullet_p(doc, "ROC-AUC (Target > 0.88), PR-AUC (Target > 0.82), Sensitivity / Recall (Target > 0.85), Specificity (Target > 0.85), F1-Score, Brier Calibration Score (< 0.10).", "ML Performance Metrics: ")

    # -------------------------------------------------------------------------
    # 19. DEPLOYMENT REQUIREMENTS
    # -------------------------------------------------------------------------
    add_heading_1(doc, "19. DEPLOYMENT ARCHITECTURE")
    add_body_p(doc, "Deployed via containerized Docker Compose services: (1) `telemed-frontend` (Vite / Nginx), (2) `telemed-backend` (FastAPI / Uvicorn), (3) `telemed-postgres` (PostgreSQL 17), and (4) `telemed-redis` (Cache / Task Queue).")

    # -------------------------------------------------------------------------
    # 20. CONSTRAINTS AND LIMITATIONS
    # -------------------------------------------------------------------------
    add_heading_1(doc, "20. CONSTRAINTS AND LIMITATIONS")
    add_bullet_p(doc, "ML prediction quality is bounded by input feature accuracy and missing modality completeness.", "Input Data Quality: ")
    add_bullet_p(doc, "Microbiome distributions vary across geographical populations and sequencing platforms.", "Microbiome Population Variance: ")
    add_bullet_p(doc, "The platform is an academic research prototype and does not hold FDA/CE clinical device clearance.", "Regulatory Status: ")

    # -------------------------------------------------------------------------
    # 21. FUTURE ENHANCEMENTS
    # -------------------------------------------------------------------------
    add_heading_1(doc, "21. FUTURE ENHANCEMENTS")
    add_bullet_p(doc, "Conduct prospective clinical validation studies with partner academic medical centers.", "Clinical Validation: ")
    add_bullet_p(doc, "Integrate FHIR / HL7 standards for automated Electronic Health Record (EHR) synchronization.", "EHR Integration: ")
    add_bullet_p(doc, "Incorporate privacy-preserving Federated Learning across distributed clinical nodes.", "Federated Learning: ")

    # -------------------------------------------------------------------------
    # 22. ACCEPTANCE CRITERIA
    # -------------------------------------------------------------------------
    add_heading_1(doc, "22. ACCEPTANCE CRITERIA")
    add_body_p(doc, "The TeleMed project shall be deemed complete upon satisfying the following 23 acceptance criteria:")

    acc_headers = ["Criteria ID", "Acceptance Verification Item", "Validation Method", "Pass / Fail Condition"]
    acc_data = [
        ["AC-01", "User Registration & Authentication", "Automated API Test", "Successful user registration & JWT login."],
        ["AC-02", "Role-Based Access Control", "Security Test", "Patients blocked from admin/doctor routes."],
        ["AC-03", "Patient Profile Creation", "Functional Test", "BMI calculated correctly from height/weight."],
        ["AC-04", "Clinical Lab Data Ingestion", "Unit Test", "18 clinical features validated against bounds."],
        ["AC-05", "Wearable CSV Data Ingestion", "Parser Test", "15 wearable/CGM features parsed accurately."],
        ["AC-06", "Microbiome Data Ingestion", "Parser Test", "40 taxa & 9 indices validated (abundance sum = 100%)."],
        ["AC-07", "Physiological Bounds Validation", "Input Test", "Invalid inputs rejected with descriptive HTTP 400."],
        ["AC-08", "Dynamic Modality Detection", "Pipeline Test", "System correctly identifies active input modalities."],
        ["AC-09", "Clinical Expert Model Execution", "ML Test", "CatBoost Clinical Expert produces valid risk scores."],
        ["AC-10", "Wearable Expert Model Execution", "ML Test", "XGBoost Wearable Expert produces valid risk scores."],
        ["AC-11", "Gut Expert Model Execution", "ML Test", "CatBoost Gut Expert produces valid risk scores."],
        ["AC-12", "7-Pathway Multimodal Fusion", "Integration Test", "Stacking Meta-Learner routes active pathway cleanly."],
        ["AC-13", "5-Disease Risk Predictions", "Inference Test", "Outputs valid probabilities for all 5 target conditions."],
        ["AC-14", "Risk Category Thresholding", "UI Test", "Scores mapped correctly to LOW, MODERATE, HIGH."],
        ["AC-15", "TreeSHAP Feature Attributions", "XAI Test", "Calculates exact Shapley values for top features."],
        ["AC-16", "PostgreSQL Data Persistence", "DB Test", "Assessment snapshots stored completely in DB."],
        ["AC-17", "Medical RAG Evidence Search", "RAG Test", "FAISS vector store returns top-3 clinical guidelines."],
        ["AC-18", "Generative AI Report Synthesis", "LLM Test", "Synthesizes structured patient & doctor briefings."],
        ["AC-19", "AI Safety Guardrail Enforcement", "Safety Test", "Rejects reports containing diagnostic assertions."],
        ["AC-20", "PDF Report Generation", "Render Test", "Generates styled, downloadable PDF assessment report."],
        ["AC-21", "Assessment History Retrieval", "UI Test", "Displays chronological history timeline."],
        ["AC-22", "Doctor Review Workflow", "Portal Test", "Verified doctor inspects assessment & adds review note."],
        ["AC-23", "Docker Container Deployment", "Ops Test", "Entire stack launches cleanly via `docker compose up`."]
    ]
    t_acc = doc.add_table(rows=1, cols=4)
    style_table(t_acc, [1.0, 2.2, 1.5, 2.0], acc_headers, acc_data)

    # -------------------------------------------------------------------------
    # 23. CONCLUSION
    # -------------------------------------------------------------------------
    add_heading_1(doc, "23. CONCLUSION")
    add_body_p(doc, 
        "This Software Requirements Specification establishes a rigorous, comprehensive blueprint for the TeleMed platform. "
        "By synthesizing clinical biomarkers, continuous wearable sensor telemetry, and gut microbiome multi-omics into an adaptive "
        "stacking fusion architecture—complemented by TreeSHAP explainability and Medical RAG report generation—TeleMed provides a state-of-the-art "
        "foundation for AI-assisted digital health decision support. This document fulfills all academic requirements for III/I B.Tech Computer Science "
        "and Engineering evaluation."
    )

    # Save Document
    doc.save(DOCX_OUTPUT_PATH)
    print(f"\nSUCCESS: TeleMed SRS Document created successfully at:\n{DOCX_OUTPUT_PATH}")

if __name__ == "__main__":
    create_srs_document()
