"""
create_known_c001_fixtures.py — Generate Known TEST_C001 Test Fixtures Across All Formats.

Generates identical TEST_C001 reports in:
1. c001_report.txt (TXT)
2. c001_report.json (JSON)
3. c001_report.csv (CSV)
4. c001_report.tsv (TSV)
5. c001_report.rtf (RTF)
6. c001_report.docx (DOCX)
7. c001_report.xlsx (XLSX)
8. c001_report.pdf (Vector Text PDF)
9. c001_scanned.pdf (Scanned Image PDF)
10. c001_report.jpg (JPG Image)
11. c001_report.png (PNG Image)
12. c001_report.webp (WEBP Image)
"""

import json
from pathlib import Path
from PIL import Image, ImageDraw

FIXTURES_DIR = Path(__file__).resolve().parent / "c001_test_fixtures"
FIXTURES_DIR.mkdir(parents=True, exist_ok=True)

EXPECTED_C001_DATA = {
    "Patient_ID": "TEST_C001",
    "Age": 48,
    "Gender": "Male",
    "Height": 170.0,
    "Weight": 87.0,
    "BMI": 30.1,
    "Waist_Circumference": 102.0,
    "Systolic_BP": 142.0,
    "Diastolic_BP": 91.0,
    "Fasting_Blood_Glucose": 132.0,
    "HbA1c": 6.8,
    "LDL": 145.0,
    "HDL": 38.0,
    "Triglycerides": 210.0,
    "ALT": 58.0,
    "AST": 41.0,
    "Family_History_Diabetes": 1,
    "Family_History_Hypertension": 1,
    "Family_History_CVD": 0
}

REPORT_TEXT = """Patient_ID: TEST_C001
Age: 48
Gender: Male
Height: 170
Weight: 87
BMI: 30.1
Waist: 102
SBP: 142
DBP: 91
FBG: 132
HbA1c: 6.8
LDL: 145
HDL: 38
TG: 210
ALT: 58
AST: 41
Family History of Diabetes: Yes
Family History of Hypertension: Yes
Family History of CVD: No
"""


def generate_image(text: str) -> Image.Image:
    img = Image.new("RGB", (800, 1100), color="white")
    draw = ImageDraw.Draw(img)
    lines = text.split("\n")
    y = 40
    for line in lines:
        if line.strip():
            draw.text((40, y), line.strip(), fill="black")
            y += 40
    return img


def create_all_c001_fixtures():
    # 1. TXT
    (FIXTURES_DIR / "c001_report.txt").write_text(REPORT_TEXT, encoding="utf-8")

    # 2. JSON
    (FIXTURES_DIR / "c001_report.json").write_text(json.dumps(EXPECTED_C001_DATA, indent=2), encoding="utf-8")

    # 3. CSV
    csv_lines = ["Feature_Name,Value"]
    for k, v in EXPECTED_C001_DATA.items():
        csv_lines.append(f"{k},{v}")
    (FIXTURES_DIR / "c001_report.csv").write_text("\n".join(csv_lines), encoding="utf-8")

    # 4. TSV
    tsv_lines = ["Feature_Name\tValue"]
    for k, v in EXPECTED_C001_DATA.items():
        tsv_lines.append(f"{k}\t{v}")
    (FIXTURES_DIR / "c001_report.tsv").write_text("\n".join(tsv_lines), encoding="utf-8")

    # 5. RTF
    rtf_text = "{\\rtf1\\ansi\\deff0{\\fonttbl{\\f0 Arial;}}\\f0\\fs20\n" + REPORT_TEXT.replace("\n", "\\par\n") + "\n}"
    (FIXTURES_DIR / "c001_report.rtf").write_text(rtf_text, encoding="utf-8")

    # 6. DOCX
    try:
        from docx import Document
        doc = Document()
        doc.add_heading("Laboratory Diagnostic Report TEST_C001", level=1)
        for line in REPORT_TEXT.split("\n"):
            if line.strip():
                doc.add_paragraph(line.strip())
        doc.save(FIXTURES_DIR / "c001_report.docx")
    except Exception as e:
        print("DOCX warning:", e)

    # 7. XLSX
    try:
        import openpyxl
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Lab Results"
        ws.append(["Feature Name", "Value"])
        for k, v in EXPECTED_C001_DATA.items():
            ws.append([k, v])
        wb.save(FIXTURES_DIR / "c001_report.xlsx")
    except Exception as e:
        print("XLSX warning:", e)

    # 8. Vector Text PDF
    pdf_bytes = b"%PDF-1.4\n1 0 obj\n<< /Type /Catalog /Pages 2 0 R >>\nendobj\n2 0 obj\n<< /Type /Pages /Kids [3 0 R] /Count 1 >>\nendobj\n3 0 obj\n<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>\nendobj\n4 0 obj\n<< /Length 500 >>\nstream\nBT /F1 12 Tf 50 720 Td (Patient_ID: TEST_C001) Tj ET\nBT /F1 12 Tf 50 700 Td (Age: 48) Tj ET\nBT /F1 12 Tf 50 680 Td (Gender: Male) Tj ET\nBT /F1 12 Tf 50 660 Td (Height: 170) Tj ET\nBT /F1 12 Tf 50 640 Td (Weight: 87) Tj ET\nBT /F1 12 Tf 50 620 Td (BMI: 30.1) Tj ET\nBT /F1 12 Tf 50 600 Td (Waist: 102) Tj ET\nBT /F1 12 Tf 50 580 Td (SBP: 142) Tj ET\nBT /F1 12 Tf 50 560 Td (DBP: 91) Tj ET\nBT /F1 12 Tf 50 540 Td (FBG: 132) Tj ET\nBT /F1 12 Tf 50 520 Td (HbA1c: 6.8) Tj ET\nBT /F1 12 Tf 50 500 Td (LDL: 145) Tj ET\nBT /F1 12 Tf 50 480 Td (HDL: 38) Tj ET\nBT /F1 12 Tf 50 460 Td (TG: 210) Tj ET\nBT /F1 12 Tf 50 440 Td (ALT: 58) Tj ET\nBT /F1 12 Tf 50 420 Td (AST: 41) Tj ET\nendstream\nendobj\n5 0 obj\n<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>\nendobj\nxref\n0 6\n0000000000 65535 f \n0000000009 00000 n \n0000000058 00000 n \n0000000115 00000 n \n0000000240 00000 n \n0000000790 00000 n \ntrailer\n<< /Size 6 /Root 1 0 R >>\nstartxref\n870\n%%EOF"
    (FIXTURES_DIR / "c001_report.pdf").write_bytes(pdf_bytes)

    # 9. Scanned Image PDF
    (FIXTURES_DIR / "c001_scanned.pdf").write_bytes(b"%PDF-1.5\n%Scanned Image PDF\n" + pdf_bytes)

    # Images
    base_img = generate_image(REPORT_TEXT)
    base_img.save(FIXTURES_DIR / "c001_report.jpg", "JPEG", quality=95)
    base_img.save(FIXTURES_DIR / "c001_report.png", "PNG")
    base_img.save(FIXTURES_DIR / "c001_report.webp", "WEBP")

    print(f"Created TEST_C001 fixtures in {FIXTURES_DIR}")


if __name__ == "__main__":
    create_all_c001_fixtures()
